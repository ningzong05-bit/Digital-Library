from pathlib import Path
import sys

from docx import Document


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [f"# {path.name}\n\n", "## Paragraphs\n\n"]
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph.text.strip()
        if text:
            parts.append(f"{index}. {text}\n")
    parts.append("\n## Tables\n\n")
    for table_index, table in enumerate(doc.tables, 1):
        parts.append(f"### Table {table_index}\n\n")
        for row in table.rows:
            cells = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |\n")
        parts.append("\n")
    return "".join(parts)


def main() -> None:
    output = Path(sys.argv[1])
    inputs = [Path(arg) for arg in sys.argv[2:]]
    content = []
    for path in inputs:
        content.append(extract_docx(path))
        content.append("\n---\n\n")
    output.write_text("".join(content), encoding="utf-8")
    print(output)


if __name__ == "__main__":
    main()
